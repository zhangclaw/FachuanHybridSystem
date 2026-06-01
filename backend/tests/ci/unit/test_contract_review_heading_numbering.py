from __future__ import annotations

import importlib.util
import sys
import types
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import pytest
from docx import Document


def _remove_numbering_part(source: str, target: str) -> None:
    """Write a copy of `source` without the optional Word numbering part."""
    with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/numbering.xml":
                continue

            if item.filename == "word/_rels/document.xml.rels":
                root = ElementTree.fromstring(data)
                for rel in list(root):
                    if rel.attrib.get("Type", "").endswith("/numbering"):
                        root.remove(rel)
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)

            if item.filename == "[Content_Types].xml":
                root = ElementTree.fromstring(data)
                for override in list(root):
                    if override.attrib.get("PartName") == "/word/numbering.xml":
                        root.remove(override)
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)

            zout.writestr(item, data)


def _load_heading_numbering_class():
    """Load the module without importing the Django-heavy services package."""
    module_path = (
        Path(__file__).resolve().parents[3]
        / "apps"
        / "contract_review"
        / "services"
        / "extraction"
        / "heading_numbering.py"
    )

    module_names = ["apps", "apps.core", "apps.core.llm", "apps.core.llm.service"]
    original_modules = {name: sys.modules.get(name) for name in module_names}

    apps_module = types.ModuleType("apps")
    core_module = types.ModuleType("apps.core")
    llm_module = types.ModuleType("apps.core.llm")
    service_module = types.ModuleType("apps.core.llm.service")
    service_module.LLMService = type("LLMService", (), {})

    sys.modules.update(
        {
            "apps": apps_module,
            "apps.core": core_module,
            "apps.core.llm": llm_module,
            "apps.core.llm.service": service_module,
        }
    )
    try:
        spec = importlib.util.spec_from_file_location("heading_numbering_under_test", module_path)
        assert spec and spec.loader
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module.HeadingNumbering
    finally:
        for name, module in original_modules.items():
            if module is None:
                sys.modules.pop(name, None)
            else:
                sys.modules[name] = module


@pytest.mark.unit
def test_get_or_add_numbering_part_creates_missing_part(tmp_path) -> None:
    original_path = tmp_path / "original.docx"
    missing_numbering_path = tmp_path / "missing-numbering.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Section heading")
    doc.add_paragraph("Subsection heading")
    doc.save(original_path)
    _remove_numbering_part(str(original_path), str(missing_numbering_path))

    broken_doc = Document(missing_numbering_path)
    with pytest.raises(NotImplementedError):
        _ = broken_doc.part.numbering_part

    service = _load_heading_numbering_class()()
    numbering_part = service._get_or_add_numbering_part(broken_doc)
    abstract_id = service._create_abstract_num(broken_doc)
    num_id = service._create_num_ref(numbering_part.element, abstract_id)
    service._apply_num_to_paragraphs(broken_doc, [(0, 0), (1, 1)], num_id)
    broken_doc.save(output_path)

    with zipfile.ZipFile(output_path, "r") as zf:
        assert any(name.startswith("word/numbering") for name in zf.namelist())
        assert "relationships/numbering" in zf.read("word/_rels/document.xml.rels").decode("utf-8")
