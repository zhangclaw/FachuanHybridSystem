"""Tests for contract_review.services.formatting.docx_revision_tool."""

from __future__ import annotations

from unittest.mock import MagicMock

import pytest

from apps.contract_review.services.formatting.docx_revision_tool import (
    DocxRevisionTool,
    _create_del,
    _create_ins,
    _make_run,
    _next_rev_id,
)


class TestNextRevId:
    def test_returns_string(self):
        result = _next_rev_id()
        assert isinstance(result, str)

    def test_increments(self):
        id1 = _next_rev_id()
        id2 = _next_rev_id()
        assert int(id2) > int(id1)


class TestMakeRun:
    def test_creates_run_element(self):
        from lxml import etree

        source = MagicMock()
        source._element = MagicMock()
        source._element.find.return_value = None
        result = _make_run("text", source)
        assert result is not None

    def test_with_del_text_tag(self):
        source = MagicMock()
        source._element = MagicMock()
        source._element.find.return_value = None
        result = _make_run("deleted", source, tag="w:delText")
        assert result is not None


class TestCreateDel:
    def test_creates_del_element(self):
        from lxml import etree

        source = MagicMock()
        source._element = MagicMock()
        source._element.find.return_value = None
        result = _create_del("old text", "author", "2024-01-01T00:00:00Z", source)
        assert result is not None


class TestCreateIns:
    def test_creates_ins_element(self):
        source = MagicMock()
        source._element = MagicMock()
        source._element.find.return_value = None
        result = _create_ins("new text", "author", "2024-01-01T00:00:00Z", source)
        assert result is not None


class TestDocxRevisionTool:
    def test_enable_track_changes(self):
        doc = MagicMock()
        settings_elem = MagicMock()
        doc.settings.element = settings_elem
        settings_elem.findall.return_value = []
        DocxRevisionTool.enable_track_changes(doc)
        settings_elem.append.assert_called_once()

    def test_apply_revision_not_found(self):
        tool = DocxRevisionTool()
        para = MagicMock()
        para.text = "no match here"
        para.runs = [MagicMock(text="no match here")]
        result = tool.apply_revision(para, "original", "replacement")
        assert result is False

    def test_apply_revision_single_run_match(self):
        """Test that single run match finds the text and returns True."""
        tool = DocxRevisionTool()
        # Create a real docx paragraph to test with
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("prefix original suffix")
        result = tool.apply_revision(para, "original", "replacement")
        assert result is True
        # Verify del and ins elements were inserted
        xml_str = para._element.xml
        assert "w:del" in xml_str
        assert "w:ins" in xml_str

    def test_apply_revision_empty_runs(self):
        tool = DocxRevisionTool()
        para = MagicMock()
        para.text = ""
        para.runs = []
        result = tool.apply_revision(para, "original", "replacement")
        assert result is False

    def test_apply_revision_cross_run(self):
        """Test cross-run matching with real docx elements."""
        tool = DocxRevisionTool()
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        # Add text split across runs
        run1 = para.add_run("hello ")
        run2 = para.add_run("world")
        result = tool.apply_revision(para, "hello world", "new text")
        assert result is True
        xml_str = para._element.xml
        assert "w:del" in xml_str
        assert "w:ins" in xml_str
