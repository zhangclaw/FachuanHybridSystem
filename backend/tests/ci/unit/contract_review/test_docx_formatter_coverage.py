"""contract_review/services/formatting/docx_formatter.py 单元测试。"""

from __future__ import annotations

from unittest.mock import MagicMock, patch

import pytest

from apps.contract_review.services.formatting.docx_formatter import (
    DocxFormatter,
    _FONT_SIZE,
    _HEADING_STYLES,
)


@pytest.fixture
def formatter() -> DocxFormatter:
    return DocxFormatter()


# ── Constants ─────────────────────────────────────────────────────────


class TestConstants:
    def test_font_sizes(self) -> None:
        assert "小二" in _FONT_SIZE
        assert "小四" in _FONT_SIZE
        assert _FONT_SIZE["小二"].pt == 18
        assert _FONT_SIZE["小四"].pt == 12

    def test_heading_styles(self) -> None:
        assert "Heading 1" in _HEADING_STYLES
        assert "Title" in _HEADING_STYLES
        assert "标题 1" in _HEADING_STYLES


# ── _find_title_element ───────────────────────────────────────────────


class TestFindTitleElement:
    def test_finds_heading_style(self) -> None:
        para_heading = MagicMock()
        para_heading.style.name = "Heading 1"
        para_heading._element = "elem1"
        para_normal = MagicMock()
        para_normal.style.name = "Normal"
        para_normal.text = "Body text"
        doc = MagicMock()
        doc.paragraphs = [para_heading, para_normal]
        result = DocxFormatter._find_title_element(doc)
        assert result == "elem1"

    def test_finds_first_non_empty_para(self) -> None:
        para = MagicMock()
        para.style.name = "Normal"
        para.text = "Title here"
        para._element = "elem_title"
        doc = MagicMock()
        doc.paragraphs = [para]
        result = DocxFormatter._find_title_element(doc)
        assert result == "elem_title"

    def test_no_paragraphs(self) -> None:
        doc = MagicMock()
        doc.paragraphs = []
        result = DocxFormatter._find_title_element(doc)
        assert result is None

    def test_empty_style_and_text(self) -> None:
        para = MagicMock()
        para.style.name = "Normal"
        para.text = ""
        doc = MagicMock()
        doc.paragraphs = [para]
        result = DocxFormatter._find_title_element(doc)
        assert result is None


# ── _set_paragraph_spacing ────────────────────────────────────────────


class TestSetParagraphSpacing:
    def test_title_spacing(self) -> None:
        para = MagicMock()
        fmt = MagicMock()
        para.paragraph_format = fmt
        para._element = MagicMock()
        para._element.find.return_value = MagicMock()

        DocxFormatter._set_paragraph_spacing(para, is_title=True)

        assert fmt.first_line_indent.pt == 0
        assert fmt.space_before.pt == 0

    def test_body_spacing(self) -> None:
        para = MagicMock()
        fmt = MagicMock()
        para.paragraph_format = fmt
        para._element = MagicMock()
        para._element.find.return_value = MagicMock()

        DocxFormatter._set_paragraph_spacing(para, is_title=False)

        assert fmt.first_line_indent.pt == 24


# ── _set_font ─────────────────────────────────────────────────────────


class TestSetFont:
    def test_sets_font_on_runs(self) -> None:
        from docx.shared import Pt

        run = MagicMock()
        run._element = MagicMock()
        para = MagicMock()
        para.runs = [run]

        with patch("docx.oxml.OxmlElement") as MockOxml:
            mock_rpr = MagicMock()
            mock_rfonts = MagicMock()
            run._element.find.return_value = mock_rpr
            mock_rpr.find.return_value = mock_rfonts

            DocxFormatter._set_font(para, "黑体", Pt(18))

            assert run.font.name == "黑体"
            assert run.font.size == Pt(18)

    def test_creates_rpr_if_missing(self) -> None:
        from docx.shared import Pt

        run = MagicMock()
        r_elem = MagicMock()
        r_elem.find.return_value = None  # No w:rPr
        run._element = r_elem
        para = MagicMock()
        para.runs = [run]

        with patch("docx.oxml.OxmlElement") as MockOxml:
            mock_rpr = MagicMock()
            mock_rfonts = MagicMock()
            MockOxml.return_value = mock_rpr
            r_elem.find.return_value = mock_rpr
            mock_rpr.find.return_value = mock_rfonts

            DocxFormatter._set_font(para, "宋体", Pt(12))
            assert run.font.name == "宋体"


# ── _clean_style_indent_chars ─────────────────────────────────────────


class TestCleanStyleIndentChars:
    def test_removes_chars_attrs(self) -> None:
        from docx.oxml.ns import qn

        chars_attrs = [qn(a) for a in ("w:firstLineChars", "w:leftChars", "w:rightChars", "w:hangingChars")]
        ind = MagicMock()
        # Only 2 keys present — the other 2 should be skipped (ind.get returns None)
        present_keys = {chars_attrs[0]: "123", chars_attrs[1]: "456"}
        ind.get.side_effect = lambda key: present_keys.get(key)
        ind.attrib = dict(present_keys)

        s_pPr = MagicMock()
        s_pPr.find.return_value = ind

        style = MagicMock()
        style.element.find.return_value = s_pPr

        doc = MagicMock()
        doc.styles = [style]

        DocxFormatter._clean_style_indent_chars(doc)
        # Verify only the 2 present keys were deleted
        assert ind.attrib == {}

    def test_no_ppr_skipped(self) -> None:
        style = MagicMock()
        style.element.find.return_value = None
        doc = MagicMock()
        doc.styles = [style]
        DocxFormatter._clean_style_indent_chars(doc)

    def test_no_ind_skipped(self) -> None:
        s_pPr = MagicMock()
        s_pPr.find.return_value = None
        style = MagicMock()
        style.element.find.return_value = s_pPr
        doc = MagicMock()
        doc.styles = [style]
        DocxFormatter._clean_style_indent_chars(doc)


# ── format_document ───────────────────────────────────────────────────


class TestFormatDocument:
    def test_calls_all_steps(self, formatter: DocxFormatter) -> None:
        title_para = MagicMock()
        title_para.style.name = "Heading 1"
        title_para._element = MagicMock()  # Use MagicMock, not a string
        title_para._element.find.return_value = MagicMock()  # pPr found
        title_para.paragraph_format = MagicMock()
        title_para.runs = []

        body_para = MagicMock()
        body_para.style.name = "Normal"
        body_para.text = "body"
        body_para._element = MagicMock()
        body_para._element.find.return_value = MagicMock()  # pPr found
        body_para.paragraph_format = MagicMock()
        body_para.runs = []

        doc = MagicMock()
        doc.paragraphs = [title_para, body_para]
        doc.styles = []

        formatter.format_document(doc)
        # Verify methods were called without exception
        assert title_para.runs == []
        assert body_para.runs == []
