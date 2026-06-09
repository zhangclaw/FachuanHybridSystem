"""Tests for workbench.services.doc_extractor."""

from __future__ import annotations

from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

from apps.workbench.services.doc_extractor import DocTextExtractor, MAX_TEXT_LENGTH


class TestDocTextExtractorInit:
    def test_init(self):
        ext = DocTextExtractor()
        assert ext._batch_converted == {}
        assert ext._batch_temp_dir is None
        assert ext._single_temp_dirs == []


class TestDocTextExtractorExtractText:
    def test_file_not_found(self):
        ext = DocTextExtractor()
        with pytest.raises(FileNotFoundError, match="文件不存在"):
            ext.extract_text("/nonexistent/file.docx")

    def test_unsupported_format(self, tmp_path):
        ext = DocTextExtractor()
        f = tmp_path / "test.xyz"
        f.write_bytes(b"data")
        with pytest.raises(ValueError, match="不支持的文件格式"):
            ext.extract_text(str(f))

    def test_txt_extraction(self, tmp_path):
        ext = DocTextExtractor()
        f = tmp_path / "test.txt"
        f.write_text("Hello World", encoding="utf-8")
        result = ext.extract_text(str(f))
        assert result == "Hello World"

    def test_txt_truncation(self, tmp_path):
        ext = DocTextExtractor()
        f = tmp_path / "large.txt"
        f.write_text("x" * (MAX_TEXT_LENGTH + 1000), encoding="utf-8")
        result = ext.extract_text(str(f))
        assert len(result) <= MAX_TEXT_LENGTH

    def test_docx_extraction(self, tmp_path):
        ext = DocTextExtractor()
        docx_file = tmp_path / "test.docx"
        # Create a real minimal docx file
        from docx import Document

        doc = Document()
        doc.add_paragraph("Para 1")
        doc.add_paragraph("")
        doc.add_paragraph("Para 3")
        doc.save(str(docx_file))
        result = ext.extract_text(str(docx_file))
        assert "Para 1" in result
        assert "Para 3" in result

    def test_doc_uses_cache(self, tmp_path):
        ext = DocTextExtractor()
        # Create a real docx file for the cache
        from docx import Document

        cached_path = str(tmp_path / "cached.docx")
        doc = Document()
        doc.add_paragraph("cached content")
        doc.save(cached_path)
        # Create a real .doc file that the path check will find
        doc_file = tmp_path / "test.doc"
        doc_file.write_bytes(b"fake doc")
        ext._batch_converted[str(doc_file)] = cached_path
        result = ext.extract_text(str(doc_file))
        assert "cached content" in result


class TestDocTextExtractorResolveDocxPath:
    def test_docx_file(self, tmp_path):
        ext = DocTextExtractor()
        f = tmp_path / "test.docx"
        f.write_bytes(b"data")
        path, cleanup = ext._resolve_docx_path(str(f))
        assert path == str(f)
        assert cleanup is False

    def test_nonexistent_file(self):
        ext = DocTextExtractor()
        path, cleanup = ext._resolve_docx_path("/nonexistent.docx")
        assert path is None
        assert cleanup is False

    def test_txt_returns_none(self, tmp_path):
        ext = DocTextExtractor()
        f = tmp_path / "test.txt"
        f.write_text("hi", encoding="utf-8")
        path, cleanup = ext._resolve_docx_path(str(f))
        assert path is None
        assert cleanup is False

    def test_doc_uses_cache(self, tmp_path):
        ext = DocTextExtractor()
        cached = tmp_path / "cached.docx"
        # Create a real docx file
        from docx import Document

        doc = Document()
        doc.add_paragraph("test")
        doc.save(str(cached))
        # Create a real .doc file
        doc_file = tmp_path / "test.doc"
        doc_file.write_bytes(b"fake doc")
        ext._batch_converted[str(doc_file)] = str(cached)
        path, cleanup = ext._resolve_docx_path(str(doc_file))
        assert path == str(cached)
        assert cleanup is False


class TestDocTextExtractorBatchConvert:
    def test_empty_paths(self):
        ext = DocTextExtractor()
        result = ext.batch_convert_doc_to_docx([])
        assert result == {}

    def test_batch_convert_calls_engine(self, tmp_path):
        ext = DocTextExtractor()
        with patch("apps.workbench.services.doc_extractor.engine_batch_convert") as mock_engine:
            mock_engine.return_value = {"/a.doc": "/b/a.docx"}
            result = ext.batch_convert_doc_to_docx(["/a.doc"], str(tmp_path))
            assert "/a.doc" in result
            assert ext._batch_converted["/a.doc"] == "/b/a.docx"


class TestDocTextExtractorCleanup:
    def test_cleanup_clears_state(self):
        ext = DocTextExtractor()
        ext._batch_converted["a"] = "b"
        ext._single_temp_dirs.append("/tmp/test")
        ext.cleanup()
        assert ext._batch_converted == {}
        assert ext._single_temp_dirs == []

    def test_cleanup_with_temp_dir(self, tmp_path):
        ext = DocTextExtractor()
        temp_dir = tmp_path / "temp"
        temp_dir.mkdir()
        ext._batch_temp_dir = str(temp_dir)
        ext.cleanup()
        assert ext._batch_temp_dir is None


class TestDocTextExtractorExtractDocMetadata:
    def test_empty_metadata_for_nonexistent(self):
        ext = DocTextExtractor()
        result = ext.extract_doc_metadata("/nonexistent.docx")
        assert result["case_number"] is None
        assert result["judge"] is None
        assert result["clerk"] is None
