"""文档处理模块测试 - 真实执行代码。"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument

from apps.automation.services.document.document_processing import (
    DocumentExtraction,
    _apply_pdf_limits,
    extract_docx_text,
    extract_document_content,
    get_doc_config,
)


class TestGetDocConfig:
    """测试文档配置获取。"""

    def test_returns_dict(self) -> None:
        config = get_doc_config()
        assert isinstance(config, dict)

    def test_has_required_keys(self) -> None:
        config = get_doc_config()
        assert "DEFAULT_TEXT_LIMIT" in config
        assert "DEFAULT_PREVIEW_PAGE" in config
        assert "MAX_TEXT_LIMIT" in config
        assert "MAX_PREVIEW_PAGES" in config

    def test_default_values(self) -> None:
        config = get_doc_config()
        assert config["DEFAULT_TEXT_LIMIT"] == 1500
        assert config["DEFAULT_PREVIEW_PAGE"] == 1
        assert config["MAX_TEXT_LIMIT"] == 10000
        assert config["MAX_PREVIEW_PAGES"] == 5


class TestApplyPdfLimits:
    """测试 PDF 限制参数应用。"""

    def test_with_explicit_values(self) -> None:
        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(500, 2, config)
        assert lim == 500
        assert page == 2

    def test_with_none_values_uses_defaults(self) -> None:
        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(None, None, config)
        assert lim == 1500
        assert page == 1

    def test_clamps_to_max(self) -> None:
        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(20000, 10, config)
        assert lim == 10000
        assert page == 5

    def test_mixed_values(self) -> None:
        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(3000, None, config)
        assert lim == 3000
        assert page == 1


class TestExtractDocxText:
    """测试 DOCX 文本提取。"""

    def test_extract_simple_docx(self, tmp_path: Path) -> None:
        """提取简单 DOCX 文本。"""
        doc = DocxDocument()
        doc.add_paragraph("这是第一段")
        doc.add_paragraph("这是第二段")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        text = extract_docx_text(str(f))
        assert "这是第一段" in text
        assert "这是第二段" in text

    def test_extract_with_limit(self, tmp_path: Path) -> None:
        """限制提取字符数。"""
        doc = DocxDocument()
        doc.add_paragraph("这是一段很长的文本内容用于测试限制功能")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        text = extract_docx_text(str(f), limit=5)
        assert len(text) <= 5

    def test_extract_empty_docx(self, tmp_path: Path) -> None:
        """提取空 DOCX。"""
        doc = DocxDocument()
        f = tmp_path / "test.docx"
        doc.save(str(f))

        text = extract_docx_text(str(f))
        assert text == ""

    def test_extract_with_table(self, tmp_path: Path) -> None:
        """提取包含表格的 DOCX。"""
        doc = DocxDocument()
        doc.add_paragraph("表格内容如下")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).paragraphs[0].text = "标题"
        f = tmp_path / "test.docx"
        doc.save(str(f))

        text = extract_docx_text(str(f))
        assert "表格内容如下" in text

    def test_extract_with_default_limit(self, tmp_path: Path) -> None:
        """不指定 limit 使用配置默认值。"""
        doc = DocxDocument()
        doc.add_paragraph("测试文本")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        text = extract_docx_text(str(f))
        assert "测试文本" in text


class TestExtractDocumentContent:
    """测试文档内容提取统一接口。"""

    def test_extract_docx(self, tmp_path: Path) -> None:
        """提取 DOCX 内容。"""
        doc = DocxDocument()
        doc.add_paragraph("测试内容")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        result = extract_document_content(str(f))
        assert isinstance(result, DocumentExtraction)
        assert result.kind == "docx"
        assert "测试内容" in (result.text or "")
        assert result.image_url is None

    def test_unsupported_extension(self, tmp_path: Path) -> None:
        """不支持的文件类型抛出异常。"""
        f = tmp_path / "test.xyz"
        f.write_bytes(b"test")

        with pytest.raises(ValueError, match="不支持的文件类型"):
            extract_document_content(str(f))

    def test_extract_docx_with_limit(self, tmp_path: Path) -> None:
        """DOCX 提取支持 limit 参数。"""
        doc = DocxDocument()
        doc.add_paragraph("这是一段测试文本内容")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        result = extract_document_content(str(f), limit=5)
        assert len(result.text or "") <= 5


class TestDocumentExtraction:
    """测试 DocumentExtraction 数据类。"""

    def test_creation(self) -> None:
        ext = DocumentExtraction(
            file_path="/tmp/test.pdf",
            text="提取的文本",
            image_url=None,
            kind="pdf",
        )
        assert ext.file_path == "/tmp/test.pdf"
        assert ext.text == "提取的文本"
        assert ext.image_url is None
        assert ext.kind == "pdf"

    def test_with_image_url(self) -> None:
        ext = DocumentExtraction(
            file_path="/tmp/test.pdf",
            text=None,
            image_url="/media/automation/processed/test.png",
            kind="pdf",
        )
        assert ext.text is None
        assert ext.image_url is not None
