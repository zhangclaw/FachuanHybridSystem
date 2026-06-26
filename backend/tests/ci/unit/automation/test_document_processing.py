"""
Tests for apps.automation.services.document.document_processing
"""

from __future__ import annotations

from unittest.mock import MagicMock, patch
from pathlib import Path

import pytest


class TestDocumentProcessing:
    """document_processing 模块测试"""

    def test_get_doc_config_default(self) -> None:
        from apps.automation.services.document.document_processing import get_doc_config

        config = get_doc_config()
        assert "DEFAULT_TEXT_LIMIT" in config
        assert "MAX_TEXT_LIMIT" in config
        assert config["DEFAULT_TEXT_LIMIT"] > 0

    def test_apply_pdf_limits_defaults(self) -> None:
        from apps.automation.services.document.document_processing import _apply_pdf_limits

        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(None, None, config)
        assert lim == 1500
        assert page == 1

    def test_apply_pdf_limits_custom(self) -> None:
        from apps.automation.services.document.document_processing import _apply_pdf_limits

        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(5000, 3, config)
        assert lim == 5000
        assert page == 3

    def test_apply_pdf_limits_clamped(self) -> None:
        from apps.automation.services.document.document_processing import _apply_pdf_limits

        config = {"DEFAULT_TEXT_LIMIT": 1500, "DEFAULT_PREVIEW_PAGE": 1, "MAX_TEXT_LIMIT": 10000, "MAX_PREVIEW_PAGES": 5}
        lim, page = _apply_pdf_limits(20000, 10, config)
        assert lim == 10000
        assert page == 5

    def test_extract_document_content_unsupported_format(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_document_content

        test_file = tmp_path / "test.xyz"
        test_file.write_text("content")
        with pytest.raises(ValueError, match="不支持的文件类型"):
            extract_document_content(str(test_file))

    def test_document_extraction_dataclass(self) -> None:
        from apps.automation.services.document.document_processing import DocumentExtraction

        ext = DocumentExtraction(file_path="/test.pdf", text="hello", image_url=None, kind="pdf")
        assert ext.file_path == "/test.pdf"
        assert ext.text == "hello"
        assert ext.kind == "pdf"

    def test_extract_docx_text(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_docx_text

        # Create a minimal docx
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("测试内容第一段")
            doc.add_paragraph("测试内容第二段")
            docx_path = tmp_path / "test.docx"
            doc.save(str(docx_path))
            text = extract_docx_text(str(docx_path), limit=100)
            assert "测试内容" in text
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_extract_docx_text_with_limit(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_docx_text

        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("A" * 100)
            doc.add_paragraph("B" * 100)
            docx_path = tmp_path / "test.docx"
            doc.save(str(docx_path))
            text = extract_docx_text(str(docx_path), limit=50)
            assert len(text) <= 50
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_extract_pdf_text(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_pdf_text

        try:
            import fitz
            pdf_path = tmp_path / "test.pdf"
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Hello World PDF测试")
            doc.save(str(pdf_path))
            doc.close()
            text = extract_pdf_text(str(pdf_path), limit=500)
            assert "Hello" in text or "PDF" in text
        except ImportError:
            pytest.skip("PyMuPDF not installed")


# ============================================================
# DocumentRenamer rename/filename 测试
# ============================================================


@pytest.mark.django_db
class TestDocumentRenamerFilename:
    """DocumentRenamer 生成文件名测试"""

    def test_generate_filename_long_title_truncated(self) -> None:
        from apps.automation.services.sms.document_renamer import DocumentRenamer

        renamer = DocumentRenamer.__new__(DocumentRenamer)
        renamer.title_extraction_limit = 50
        long_title = "这是一个非常非常非常非常非常非常非常长的文书标题"
        filename = renamer.generate_filename(long_title, "案件名", __import__("datetime").date(2025, 1, 1))
        assert filename.endswith(".pdf")

    def test_generate_filename_special_chars(self) -> None:
        from apps.automation.services.sms.document_renamer import DocumentRenamer

        renamer = DocumentRenamer.__new__(DocumentRenamer)
        renamer.title_extraction_limit = 50
        filename = renamer.generate_filename("判决<>书", "张三:*案件", __import__("datetime").date(2025, 6, 1))
        assert "<" not in filename
        assert ">" not in filename
        assert ":" not in filename

    def test_rename_with_fallback_file_not_found(self, tmp_path) -> None:
        from apps.automation.services.sms.document_renamer import DocumentRenamer
        from apps.core.exceptions import ValidationException

        renamer = DocumentRenamer.__new__(DocumentRenamer)
        renamer.title_extraction_limit = 50
        with pytest.raises(ValidationException):
            renamer.rename(str(tmp_path / "nonexistent.pdf"), "案件名", __import__("datetime").date(2025, 1, 1))


# ============================================================
# Extended coverage tests
# ============================================================


class TestSaveUploadedDocument:
    def test_save_file(self, tmp_path) -> None:
        from django.core.files.uploadedfile import SimpleUploadedFile
        from apps.automation.services.document.document_processing import save_uploaded_document

        upload = SimpleUploadedFile(name="test_upload.pdf", content=b"%PDF-1.4chunk2", content_type="application/pdf")

        from django.test import override_settings
        with override_settings(MEDIA_ROOT=str(tmp_path)):
            result = save_uploaded_document(upload)

        assert result.exists()
        assert result.suffix == ".pdf"
        result.unlink(missing_ok=True)

    def test_unique_filenames(self, tmp_path) -> None:
        from django.core.files.uploadedfile import SimpleUploadedFile
        from apps.automation.services.document.document_processing import save_uploaded_document

        upload1 = SimpleUploadedFile(name="same.pdf", content=b"%PDF-1.4data1", content_type="application/pdf")
        upload2 = SimpleUploadedFile(name="same.pdf", content=b"%PDF-1.4data2", content_type="application/pdf")

        from django.test import override_settings
        with override_settings(MEDIA_ROOT=str(tmp_path)):
            r1 = save_uploaded_document(upload1)
            r2 = save_uploaded_document(upload2)

        assert r1.name != r2.name
        r1.unlink(missing_ok=True)
        r2.unlink(missing_ok=True)


class TestRenderPdfPageToImage:
    def test_basic_render(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import render_pdf_page_to_image

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Test Content")
        doc.save(str(pdf_path))
        doc.close()

        with patch("apps.automation.services.document.document_processing.settings") as mock_settings:
            mock_settings.MEDIA_ROOT = str(tmp_path / "media")
            mock_settings.MEDIA_URL = "/media/"
            url = render_pdf_page_to_image(str(pdf_path), page_num=0)

        assert "/media/" in url
        assert url.endswith(".png")

    def test_negative_page_clamped(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import render_pdf_page_to_image

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()

        with patch("apps.automation.services.document.document_processing.settings") as mock_settings:
            mock_settings.MEDIA_ROOT = str(tmp_path / "media")
            mock_settings.MEDIA_URL = "/media/"
            url = render_pdf_page_to_image(str(pdf_path), page_num=-5)

        assert url  # negative clamped to 0

    def test_page_exceeds_count(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import render_pdf_page_to_image

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()

        with patch("apps.automation.services.document.document_processing.settings") as mock_settings:
            mock_settings.MEDIA_ROOT = str(tmp_path / "media")
            mock_settings.MEDIA_URL = "/media/"
            url = render_pdf_page_to_image(str(pdf_path), page_num=999)

        assert url  # clamped to last page

    def test_render_first_page_alias(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import render_pdf_first_page_to_image

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()

        with patch("apps.automation.services.document.document_processing.settings") as mock_settings:
            mock_settings.MEDIA_ROOT = str(tmp_path / "media")
            mock_settings.MEDIA_URL = "/media/"
            url = render_pdf_first_page_to_image(str(pdf_path))

        assert url


class TestExtractDocumentContentExtended:
    def test_pdf_content(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import extract_pdf_text

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello World PDF Test 12345")
        doc.save(str(pdf_path))
        doc.close()

        text = extract_pdf_text(str(pdf_path))
        assert "Hello World" in text

    def test_docx_content(self, tmp_path) -> None:
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from apps.automation.services.document.document_processing import extract_document_content

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("DOCX提取内容")
        doc.save(str(docx_path))

        result = extract_document_content(str(docx_path))
        assert result.kind == "docx"
        assert "DOCX提取内容" in result.text

    def test_image_content(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import extract_document_content

        img_path = tmp_path / "test.jpg"
        img_path.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 100)  # minimal JPEG header

        with patch(
            "apps.automation.services.document.document_processing.extract_text_from_image_with_rapidocr"
        ) as mock_ocr:
            mock_ocr.return_value = "图片OCR文字"
            result = extract_document_content(str(img_path))

        assert result.kind == "image"
        assert "图片OCR文字" in result.text


class TestOcrPdfPage:
    def test_ocr_with_valid_pdf(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import _ocr_pdf_page

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "OCR测试文字")
        doc.save(str(pdf_path))
        doc.close()

        with patch("apps.automation.services.document.document_processing.settings") as mock_settings:
            mock_settings.MEDIA_ROOT = str(tmp_path / "media")
            with patch(
                "apps.automation.services.document.document_processing.extract_text_from_image_with_rapidocr"
            ) as mock_ocr:
                mock_ocr.return_value = "OCR识别结果"
                result = _ocr_pdf_page(str(pdf_path), 1, 500)

        assert result == "OCR识别结果"

    def test_ocr_page_out_of_range(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import _ocr_pdf_page

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()

        with patch("apps.automation.services.document.document_processing.settings") as mock_settings:
            mock_settings.MEDIA_ROOT = str(tmp_path / "media")
            with patch(
                "apps.automation.services.document.document_processing.extract_text_from_image_with_rapidocr"
            ) as mock_ocr:
                mock_ocr.return_value = "text"
                result = _ocr_pdf_page(str(pdf_path), 999, 500)

        # Should still return text since page_num is clamped
        assert result == "text"

    def test_ocr_exception_returns_none(self, tmp_path) -> None:
        from apps.automation.services.document.document_processing import _ocr_pdf_page

        result = _ocr_pdf_page("/nonexistent/file.pdf", 1, 500)
        assert result is None


class TestProcessPdfExtended:
    def test_text_pdf_returns_text(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import extract_pdf_text

        pdf_path = tmp_path / "text.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Extractable Text Content")
        doc.save(str(pdf_path))
        doc.close()

        text = extract_pdf_text(str(pdf_path))
        assert "Extractable Text Content" in text

    def test_empty_pdf_falls_back_to_ocr(self, tmp_path) -> None:
        import fitz

        from apps.automation.services.document.document_processing import process_pdf

        pdf_path = tmp_path / "empty.pdf"
        doc = fitz.open()
        doc.new_page()  # Empty page - no text
        doc.save(str(pdf_path))
        doc.close()

        with patch("apps.automation.services.document.document_processing.settings") as mock_settings:
            mock_settings.MEDIA_ROOT = str(tmp_path / "media")
            mock_settings.MEDIA_URL = "/media/"
            with patch(
                "apps.automation.services.document.document_processing.extract_text_from_image_with_rapidocr"
            ) as mock_ocr:
                mock_ocr.return_value = "OCR结果"
                image_url, text = process_pdf(str(pdf_path))

        # Should get OCR text or image URL
        assert text is not None or image_url is not None
