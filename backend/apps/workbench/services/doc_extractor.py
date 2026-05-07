"""文档文本提取器

支持 .doc 和 .docx 两种格式：
- .docx: python-docx 直接解析
- .doc: LibreOffice 批量转换为 .docx，再用 python-docx 解析

已验证方案：
- mammoth 不支持 .doc（仅 ZIP/XML）
- antiword 中文乱码（GBK 不兼容）
- textutil (macOS) 输出空字符
- LibreOffice headless 中文完美，批量转换 195 文件约 20 秒
"""

from __future__ import annotations

import logging
import platform
import shutil
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_TEXT_LENGTH = 100000  # 上限保护，长文档由 tasks.py 分段处理
BATCH_CONVERT_SIZE = 25  # 每批转换的文件数


class DocTextExtractor:
    """文档文本提取器"""

    def __init__(self) -> None:
        self._batch_converted: dict[str, str] = {}  # .doc 路径 → 转换后的 .docx 路径
        self._batch_temp_dir: str | None = None  # 批量转换的临时目录
        self._single_temp_dirs: list[str] = []  # 单文件转换的临时目录

    def extract_text(self, file_path: str) -> str:
        """根据文件扩展名选择提取策略

        Args:
            file_path: 文件路径

        Returns:
            提取的文本内容（截断到 MAX_TEXT_LENGTH）

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext == ".docx":
            return self._extract_docx(file_path)
        elif ext == ".doc":
            return self._extract_doc(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    # 判决书尾部法官/书记员正则
    _JUDGE_RE = __import__("re").compile(r"审\s*判\s*(?:长|员)\s*[：:]\s*(.+)")
    _ASSESSOR_RE = __import__("re").compile(r"人民陪审员\s*[：:]\s*(.+)")
    _CLERK_RE = __import__("re").compile(r"书\s*记\s*员\s*[：:]\s*(.+)")

    def extract_doc_metadata(self, file_path: str) -> dict[str, str | None]:
        """从文档中提取元数据

        - 首部表格：案号、审理法院、裁判日期、案由
        - 尾部段落：审判长/审判员、书记员

        Returns:
            {"case_number", "court", "judgment_date", "cause", "judge", "clerk"}
        """
        docx_path, need_cleanup = self._resolve_docx_path(file_path)
        empty = {"case_number": None, "court": None, "judgment_date": None, "cause": None, "judge": None, "clerk": None}
        if not docx_path:
            return empty

        try:
            import re

            from docx import Document

            doc = Document(docx_path)
            metadata = dict(empty)

            # ── 首部表格 ──
            key_map = {"案号": "case_number", "审理法院": "court", "裁判日期": "judgment_date", "案由": "cause"}
            for table in doc.tables[:3]:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 3:
                        key = key_map.get(cells[0])
                        if key and cells[2]:
                            metadata[key] = cells[2]
                if metadata["case_number"]:
                    break

            # ── 尾部段落：法官/书记员 ──
            paragraphs = doc.paragraphs[-15:]
            for para in paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                m = self._JUDGE_RE.search(text)
                if m:
                    metadata["judge"] = m.group(1).replace(" ", "").replace("　", "")
                    continue
                m = self._CLERK_RE.search(text)
                if m:
                    metadata["clerk"] = m.group(1).replace(" ", "").replace("　", "")

            return metadata
        except Exception:
            logger.warning("提取文档元数据失败: %s", file_path, exc_info=True)
            return dict(empty)
        finally:
            if need_cleanup and docx_path != file_path:
                Path(docx_path).unlink(missing_ok=True)

    def _resolve_docx_path(self, file_path: str) -> tuple[str | None, bool]:
        """解析文件路径，返回 (docx_path, need_cleanup)"""
        path = Path(file_path)
        if not path.exists():
            return None, False

        ext = path.suffix.lower()
        if ext == ".docx":
            return file_path, False
        elif ext == ".doc":
            cached = self._batch_converted.get(file_path)
            if cached and Path(cached).exists():
                return cached, False
            try:
                return self._convert_single_doc(file_path), True
            except Exception:
                logger.warning("转换 .doc 文件失败: %s", file_path, exc_info=True)
                return None, False
        return None, False

    def _extract_docx(self, path: str) -> str:
        """用 python-docx 提取 .docx 文本"""
        from docx import Document

        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if len(text) > MAX_TEXT_LENGTH:
            logger.warning("文本截断: %s (%d -> %d 字符)", path, len(text), MAX_TEXT_LENGTH)
            text = text[:MAX_TEXT_LENGTH]
        return text

    def _extract_doc(self, path: str) -> str:
        """将 .doc 转换为 .docx 后提取文本"""
        # 优先使用批量转换的缓存结果
        cached = self._batch_converted.get(path)
        if cached and Path(cached).exists():
            return self._extract_docx(cached)

        docx_path = self._convert_single_doc(path)
        try:
            return self._extract_docx(docx_path)
        finally:
            Path(docx_path).unlink(missing_ok=True)

    def batch_convert_doc_to_docx(
        self,
        doc_paths: list[str],
        output_dir: str | None = None,
    ) -> dict[str, str]:
        """批量将 .doc 转换为 .docx

        LibreOffice 支持单次传入多个文件，JVM 启动开销被分摊。
        实测：195 个文件分 10 批约 20 秒完成。

        Args:
            doc_paths: .doc 文件路径列表
            output_dir: 输出目录（默认创建临时目录）

        Returns:
            {原始 .doc 路径: 转换后的 .docx 路径}
        """
        if not doc_paths:
            return {}

        soffice = self._find_libreoffice()
        if not soffice:
            raise RuntimeError(
                "未找到 LibreOffice，无法转换 .doc 文件。请安装 LibreOffice: https://www.libreoffice.org/"
            )

        if output_dir is None:
            output_dir = tempfile.mkdtemp(prefix="workbench_batch_")
            self._batch_temp_dir = output_dir

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        result: dict[str, str] = {}

        # 分批转换
        for i in range(0, len(doc_paths), BATCH_CONVERT_SIZE):
            batch = doc_paths[i : i + BATCH_CONVERT_SIZE]
            logger.info(
                "LibreOffice 批量转换: 第 %d-%d/%d 个文件",
                i + 1,
                min(i + BATCH_CONVERT_SIZE, len(doc_paths)),
                len(doc_paths),
            )

            cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", str(output_path)] + batch
            try:
                proc = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                if proc.returncode != 0:
                    logger.error("LibreOffice 转换失败: %s", proc.stderr)
                    # 回退到逐个转换
                    for doc_path in batch:
                        try:
                            docx_path = self._convert_single_doc(doc_path)
                            result[doc_path] = docx_path
                        except Exception as e:
                            logger.error("单文件转换失败: %s - %s", doc_path, e)
                    continue

                # 映射转换结果
                for doc_path in batch:
                    doc_name = Path(doc_path).stem + ".docx"
                    docx_path = output_path / doc_name
                    if docx_path.exists():
                        result[doc_path] = str(docx_path)
                    else:
                        logger.warning("转换后文件未找到: %s", docx_path)

            except subprocess.TimeoutExpired:
                logger.error("LibreOffice 转换超时")
                # 回退到逐个转换
                for doc_path in batch:
                    try:
                        docx_path = self._convert_single_doc(doc_path)
                        result[doc_path] = docx_path
                    except Exception as e:
                        logger.error("单文件转换失败: %s - %s", doc_path, e)

        logger.info("批量转换完成: %d/%d 成功", len(result), len(doc_paths))
        self._batch_converted.update(result)
        return result

    async def batch_convert_doc_to_docx_async(
        self,
        doc_paths: list[str],
        output_dir: str | None = None,
        parallel_batches: int = 3,
    ) -> dict[str, str]:
        """异步批量将 .doc 转换为 .docx

        用 asyncio.create_subprocess_exec 并行运行多个 LibreOffice 批次。

        Args:
            doc_paths: .doc 文件路径列表
            output_dir: 输出目录
            parallel_batches: 并行批次数

        Returns:
            {原始 .doc 路径: 转换后的 .docx 路径}
        """
        import asyncio

        if not doc_paths:
            return {}

        soffice = self._find_libreoffice()
        if not soffice:
            raise RuntimeError(
                "未找到 LibreOffice，无法转换 .doc 文件。请安装 LibreOffice: https://www.libreoffice.org/"
            )

        if output_dir is None:
            output_dir = tempfile.mkdtemp(prefix="workbench_batch_")
            self._batch_temp_dir = output_dir

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        # 分批
        batches: list[list[str]] = []
        for i in range(0, len(doc_paths), BATCH_CONVERT_SIZE):
            batches.append(doc_paths[i : i + BATCH_CONVERT_SIZE])

        result: dict[str, str] = {}

        # 并行执行批次
        for batch_group_start in range(0, len(batches), parallel_batches):
            group = batches[batch_group_start : batch_group_start + parallel_batches]
            logger.info(
                "LibreOffice 异步批量转换: 第 %d-%d 批 (共 %d 批)",
                batch_group_start + 1,
                min(batch_group_start + parallel_batches, len(batches)),
                len(batches),
            )

            async def run_batch(batch: list[str]) -> None:
                cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", str(output_path)] + batch
                proc = await asyncio.create_subprocess_exec(
                    *cmd,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE,
                )
                stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)
                if proc.returncode != 0:
                    logger.error("LibreOffice 异步转换失败: %s", stderr.decode(errors="replace"))
                    # 回退到同步逐个转换
                    for doc_path in batch:
                        try:
                            docx_path = self._convert_single_doc(doc_path)
                            result[doc_path] = docx_path
                        except Exception as e:
                            logger.error("单文件转换失败: %s - %s", doc_path, e)
                    return

                for doc_path in batch:
                    doc_name = Path(doc_path).stem + ".docx"
                    docx_path = output_path / doc_name
                    if docx_path.exists():
                        result[doc_path] = str(docx_path)
                    else:
                        logger.warning("转换后文件未找到: %s", docx_path)

            await asyncio.gather(*(run_batch(b) for b in group))

        logger.info("异步批量转换完成: %d/%d 成功", len(result), len(doc_paths))
        self._batch_converted.update(result)
        return result

    def _convert_single_doc(self, doc_path: str) -> str:
        """单个 .doc 转 .docx"""
        soffice = self._find_libreoffice()
        if not soffice:
            raise RuntimeError("未找到 LibreOffice")

        output_dir = tempfile.mkdtemp(prefix="workbench_doc_")
        self._single_temp_dirs.append(output_dir)
        cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", output_dir, doc_path]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice 转换失败: {proc.stderr}")

        doc_name = Path(doc_path).stem + ".docx"
        docx_path = Path(output_dir) / doc_name
        if not docx_path.exists():
            raise FileNotFoundError(f"转换后文件未找到: {docx_path}")

        return str(docx_path)

    @staticmethod
    def _find_libreoffice() -> str | None:
        """查找 LibreOffice 可执行路径

        复用 apps/documents/services/infrastructure/pdf_merge_utils.py 的逻辑。
        """
        # 1. PATH 中查找
        path = shutil.which("soffice") or shutil.which("libreoffice")
        if path:
            return path

        # 2. macOS 标准安装路径
        if platform.system() == "Darwin":
            mac_paths = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "/Applications/OpenOffice.app/Contents/MacOS/soffice",
            ]
            for p in mac_paths:
                if Path(p).exists():
                    return p

        # 3. Linux 标准安装路径
        if platform.system() == "Linux":
            linux_paths = [
                "/usr/bin/libreoffice",
                "/usr/bin/soffice",
                "/usr/local/bin/libreoffice",
                "/snap/bin/libreoffice",
            ]
            for p in linux_paths:
                if Path(p).exists():
                    return p

        return None

    def cleanup(self) -> None:
        """清理批量转换和单文件转换产生的临时目录"""
        if self._batch_temp_dir:
            temp_path = Path(self._batch_temp_dir)
            if temp_path.exists():
                shutil.rmtree(temp_path, ignore_errors=True)
                logger.info("已清理批量转换临时目录: %s", self._batch_temp_dir)
            self._batch_temp_dir = None
        for dir_path in self._single_temp_dirs:
            p = Path(dir_path)
            if p.exists():
                shutil.rmtree(p, ignore_errors=True)
        if self._single_temp_dirs:
            logger.info("已清理 %d 个单文件转换临时目录", len(self._single_temp_dirs))
        self._single_temp_dirs.clear()
        self._batch_converted.clear()
